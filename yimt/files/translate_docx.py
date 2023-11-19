""""DOC/DOCX file translation"""
import argparse

import os
import re
from io import BytesIO

import docx
from docx import ImagePart
from docx.document import Document
from docx.oxml import CT_Picture
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from yimt.api.utils import detect_lang

docx_progress = ""
def doc_to_docx(doc_fn, docx_fn):
    from win32com import client as wc
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(doc_fn)
    doc.SaveAs(docx_fn, 12, False, "", True, "", False, False, False, False)
    doc.Close()
    word.Quit()


def is_pic(paragraph: Paragraph):
    img = paragraph._element.xpath('.//pic:pic')
    if not img:
        return False
    else:
        return True


def handle_sections(doc, new_doc):
    """Copy page format"""
    sec = doc.sections[0]
    sec_new = new_doc.sections[0]
    sec_new.left_margin = sec.left_margin
    sec_new.right_margin = sec.right_margin
    sec_new.top_margin = sec.top_margin
    sec_new.bottom_margin = sec.bottom_margin
    sec_new.header_distance = sec.header_distance
    sec_new.footer_distance = sec.footer_distance
    sec_new.orientation = sec.orientation
    sec_new.page_height = sec.page_height
    sec_new.page_width = sec.page_width


def get_heading(p):
    s = p.style.name
    m = re.match(r"Heading (\d)", s)
    if m is None:
        return 0

    return int(m.group(1))


def handle_paragraph_txt(p, new_doc):
    """Copy and get text to be translated"""
    runs = []
    h = get_heading(p)
    # print("p.alignment：")
    # print(p.alignment)
    if p.alignment == None:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # print(p.paragraph_format.alignment)
    if h == 0:
        new_p = new_doc.add_paragraph()
        new_p.paragraph_format.alignment = p.paragraph_format.alignment
        new_p.paragraph_format.left_indent = p.paragraph_format.left_indent
        new_p.paragraph_format.right_indent = p.paragraph_format.right_indent
        new_p.paragraph_format.first_line_indent = p.paragraph_format.first_line_indent
        new_p.paragraph_format.line_spacing = p.paragraph_format.line_spacing
        new_p.paragraph_format.space_before = p.paragraph_format.space_before
        new_p.paragraph_format.space_after = p.paragraph_format.space_after
    else:
        new_p = new_doc.add_heading(level=h)
        new_p.paragraph_format.alignment = p.paragraph_format.alignment
        new_p.paragraph_format.left_indent = p.paragraph_format.left_indent
        new_p.paragraph_format.right_indent = p.paragraph_format.right_indent
        new_p.paragraph_format.first_line_indent = p.paragraph_format.first_line_indent
        new_p.paragraph_format.line_spacing = p.paragraph_format.line_spacing
        new_p.paragraph_format.space_before = p.paragraph_format.space_before
        new_p.paragraph_format.space_after = p.paragraph_format.space_after

    for r in p.runs:
        new_r = new_p.add_run()
        new_r.text = r.text
        new_r.font.size = r.font.size
        new_r.font.bold = r.font.bold
        new_r.font.italic = r.font.italic
        new_r.font.color.rgb = r.font.color.rgb
        new_r.font.highlight_color = r.font.highlight_color
        runs.append(new_r)

    return runs


def handle_paragraph_img(p, doc, new_doc):
    """Copy image"""
    img = p._element.xpath('.//pic:pic')
    img: CT_Picture = img[0]
    embed = img.xpath('.//a:blip/@r:embed')[0]
    related_part: ImagePart = doc.part.related_parts[embed]
    # image: Image = related_part.image

    new_img = new_doc.add_picture(BytesIO(related_part.blob))


def scan_doc(doc, new_doc):
    if isinstance(doc, Document):
        parent_elm = doc.element.body
    elif isinstance(doc, _Cell):
        parent_elm = doc._tc
    else:
        raise ValueError("something's not right")

    handle_sections(doc, new_doc)

    runs = []
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            p = Paragraph(child, doc)
            if is_pic(p):
                # print("Image P")
                handle_paragraph_img(p, doc, new_doc)
            else:
                # print("Text P:", p.alignment)
                runs.extend(handle_paragraph_txt(p, new_doc))
            # yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            # yield Table(child, parent)
            table = Table(child, doc)
            new_table = new_doc.add_table(len(table.rows), len(table.columns))
            for i in range(len(table.rows)):
                row = table.rows[i]
                for j in range(len(row.cells)):
                    cell = row.cells[j]
                    for paragraph in cell.paragraphs:
                        if is_pic(paragraph):
                            # print("Image P")
                            handle_paragraph_img(paragraph, doc, new_doc)
                        else:
                            # print("Text P:", paragraph.text)
                            new_table.rows[i].cells[j].text = paragraph.text
                            runs.append(new_table.rows[i].cells[j])

    return runs


def translate_docx_auto(in_fn, source_lang="auto", target_lang="zh", translation_file=None):
    paths = os.path.splitext(in_fn)
    doc_type = paths[1]
    docx_fn = in_fn
    if doc_type == ".doc":
        docx_fn = paths[0] + ".docx"
        print("Converting doc type into docx type...")
        doc_to_docx(in_fn, docx_fn)

    if translation_file is None:
        translated_fn = paths[0] + "-translated.docx"
    else:
        translated_fn = translation_file

    doc = docx.Document(docx_fn)
    translated_doc = docx.Document()
    runs = scan_doc(doc, translated_doc)

    if source_lang == "auto":
        source_lang = detect_lang(runs[0].text)

    from yimt.api.translators import Translators
    translator = Translators().get_translator(source_lang, target_lang)

    txt_list = [r.text for r in runs]
    global docx_progress
    docx_progress = ""
    batch_size = 10  # 每多少个文本更新一个进度单位
    result_list = []
    for i in range(0, len(txt_list) // batch_size + 1):
        batch = txt_list[i * batch_size: i * batch_size + batch_size]
        # print(batch) # 测试用
        result = translator.translate_list(batch)
        result_list += result
        docx_progress += "#"
        print("docx_progress:" + docx_progress)  # 测试用
    # result_list = translator.translate_list(txt_list)
    for i in range(len(runs)):
        runs[i].text = result_list[i]

    translated_doc.save(translated_fn)
    docx_progress = ""
    return translated_fn


if __name__ == "__main__":
    arg_parser = argparse.ArgumentParser("DOC File Translator")
    arg_parser.add_argument("--to_lang", type=str, default="zh", help="target language")
    arg_parser.add_argument("--input_file", type=str, required=True, help="file to be translated")
    arg_parser.add_argument("--output_file", type=str, default=None, help="translation file")
    args = arg_parser.parse_args()

    in_file = args.input_file
    out_file = args.output_file
    to_lang = args.to_lang

    translated_fn = translate_docx_auto(in_file, target_lang=to_lang, translation_file=out_file)

    import webbrowser

    webbrowser.open(in_file)
    webbrowser.open(translated_fn)

